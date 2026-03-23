import os
from pathlib import Path
from typing import Dict, List

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


CHINESE_DIGITS = "零一二三四五六七八九"
CHINESE_UNITS = ["", "十", "百", "千"]


def _set_fangsong(run, size_pt: int, bold: bool):
    run.bold = bold
    run.font.name = "仿宋"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋")
    run.font.size = Pt(size_pt)


def _int_to_chinese(num: int) -> str:
    if num == 0:
        return CHINESE_DIGITS[0]

    parts = []
    unit_pos = 0
    need_zero = False

    while num > 0:
        digit = num % 10
        if digit == 0:
            if parts and not need_zero:
                need_zero = True
        else:
            if need_zero:
                parts.append(CHINESE_DIGITS[0])
                need_zero = False
            parts.append(CHINESE_UNITS[unit_pos])
            parts.append(CHINESE_DIGITS[digit])
        unit_pos += 1
        num //= 10

    result = "".join(reversed(parts))
    result = result.replace("一十", "十", 1)
    return result


class WordReportService:
    def __init__(self, output_dir: str = "reports"):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(exist_ok=True)

    def build_project_report(
        self,
        project_name: str,
        supplier_names: List[str],
        supplier_screenshot_map: Dict[str, str],
    ) -> str:
        document = Document()

        title = document.add_paragraph()
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        title_run = title.add_run(f"<{project_name}>关联性报告")
        _set_fangsong(title_run, size_pt=22, bold=True)

        document.add_paragraph("")

        for idx, supplier_name in enumerate(supplier_names, start=1):
            name_para = document.add_paragraph()
            name_run = name_para.add_run(f"{idx}. {supplier_name}")
            _set_fangsong(name_run, size_pt=16, bold=True)

            image_path = supplier_screenshot_map.get(supplier_name)
            if image_path and os.path.exists(image_path):
                document.add_picture(image_path, width=Inches(7.0))
            else:
                missing_para = document.add_paragraph("（未生成截图）")
                missing_run = missing_para.runs[0]
                _set_fangsong(missing_run, size_pt=14, bold=False)

            document.add_paragraph("")

        supplier_count_cn = _int_to_chinese(len(supplier_names))
        conclusion_para = document.add_paragraph()
        conclusion_run = conclusion_para.add_run(
            f"结论：经核查，以上{supplier_count_cn}家单位无关联性。"
        )
        _set_fangsong(conclusion_run, size_pt=16, bold=True)

        report_path = self.output_dir / f"{project_name}关联性报告.docx"
        document.save(str(report_path))
        return str(report_path)
